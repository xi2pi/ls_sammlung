# -*- coding: utf-8 -*-
"""
@author: Christian Winkler
"""

import pandas as pd
from docx import Document
import datetime

# Beispiel-Excel-Datei einlesen
excel_datei = "src_data/ls_all.xlsx"
df = pd.read_excel(excel_datei)

# Überprüfen, ob die erforderlichen Spalten existieren
if "aktenzeichen" not in df.columns or "leitsatz" not in df.columns:
    raise ValueError("Die Spalten 'aktenzeichen' und/oder 'leitsatz' fehlen in der Excel-Datei.")

# Datum abrufen
date_today = datetime.datetime.now().strftime("%Y-%m-%d")

# Anzahl der extrahierten Leitsätze
extrahiert = df["leitsatz"].notnull().sum()

# Word-Dokument erstellen
def create_word_file(dataframe, filename="output.docx"):
    document = Document()
    # Titel
    document.add_heading("Leitsätze des BGH (X. und Xa. Senat) und des BPatG im Zeitraum 2000 bis 2024\n", level=0)

    # Absatz mit Informationen
    info_paragraph = document.add_paragraph()
    info_paragraph.add_run("Datum: " + date_today + "\n").bold = True
    info_paragraph.add_run("Anzahl der Entscheidungen: " + str(extrahiert) + "\n\n").bold = True
    info_paragraph.add_run("Quellen:\n\n")
    info_paragraph.add_run("Fobbe, S. (2024). Corpus der Entscheidungen des Bundesgerichtshofs (CE-BGH) (2024-09-25) [Data set]. Zenodo. https://doi.org/10.5281/zenodo.12814022\n\n")
    info_paragraph.add_run("Fobbe, S. (2024). Corpus der Entscheidungen des Bundespatentgerichts (CE-BPatG) (2024-07-09) [Data set]. Zenodo. https://doi.org/10.5281/zenodo.10849977")

    # Untertitel "Leitsätze"
    document.add_heading("Entscheidungen", level=1)

    # Daten aus dem DataFrame in das Word-Dokument einfügen
    for _, row in dataframe.iterrows():
        aktenzeichen = row["aktenzeichen"]
        text = row["leitsatz"]

        if pd.notna(aktenzeichen) and pd.notna(text):  # Überprüfen, ob keine leeren Werte vorliegen
            document.add_heading(str(aktenzeichen), level=2)  # Aktenzeichen als kleine Überschrift
            document.add_paragraph(str(text))  # Text unterhalb der Überschrift

    document.save(filename)

# TXT-Datei erstellen
def create_txt_file(dataframe, filename="output.txt"):
    with open(filename, "w", encoding="utf-8") as file:
        file.write("Leitsätze des BGH (X. und Xa. Senat) und des BPatG im Zeitraum 2000 bis 2024\n\n")
        file.write(f"Datum: {date_today}\n")
        file.write(f"Anzahl der Entscheidungen: {extrahiert}\n\n")
        file.write("Quellen:\n")
        file.write("Fobbe, S. (2024). Corpus der Entscheidungen des Bundesgerichtshofs (CE-BGH) (2024-09-25) [Data set]. Zenodo. https://doi.org/10.5281/zenodo.12814022\n")
        file.write("Fobbe, S. (2024). Corpus der Entscheidungen des Bundespatentgerichts (CE-BPatG) (2024-07-09) [Data set]. Zenodo. https://doi.org/10.5281/zenodo.10849977\n\n")
        file.write("Entscheidungen:\n\n")
        for _, row in dataframe.iterrows():
            aktenzeichen = row["aktenzeichen"]
            text = row["leitsatz"]
            if pd.notna(aktenzeichen) and pd.notna(text):
                file.write(f"{aktenzeichen}\n")
                file.write(f"{text}\n\n")

# Markdown-Datei erstellen
def create_md_file(dataframe, filename="output.md"):
    with open(filename, "w", encoding="utf-8") as file:
        file.write("# Leitsätze des BGH (X. und Xa. Senat) und des BPatG im Zeitraum 2000 bis 2024\n\n")
        file.write(f"**Datum**: {date_today}  \n")
        file.write(f"**Anzahl der Entscheidungen**: {extrahiert}  \n\n")
        file.write("## Quellen\n")
        file.write("Fobbe, S. (2024). Corpus der Entscheidungen des Bundesgerichtshofs (CE-BGH) (2024-09-25) [Data set]. Zenodo. https://doi.org/10.5281/zenodo.12814022  \n")
        file.write("Fobbe, S. (2024). Corpus der Entscheidungen des Bundespatentgerichts (CE-BPatG) (2024-07-09) [Data set]. Zenodo. https://doi.org/10.5281/zenodo.10849977  \n\n")
        file.write("## Entscheidungen\n\n")
        for _, row in dataframe.iterrows():
            aktenzeichen = row["aktenzeichen"]
            text = row["leitsatz"]
            if pd.notna(aktenzeichen) and pd.notna(text):
                file.write(f"### {aktenzeichen}\n")
                file.write(f"{text}\n\n")
                
                
# Dateien erstellen

datei_bezeichnung = "sammlung/sammlung_bgh_bpatg"

create_word_file(df, datei_bezeichnung+".docx")
create_txt_file(df, datei_bezeichnung+".txt")
create_md_file(df, datei_bezeichnung+".md")

print("Alle Dateien wurden erfolgreich erstellt: Word, TXT, Markdown.")
